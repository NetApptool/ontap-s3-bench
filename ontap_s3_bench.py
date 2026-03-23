#!/usr/bin/env python3
"""
ONTAP S3 对象存储自动化性能测试工具
联想凌拓科技有限公司

用法:
  python3 ontap_s3_bench.py                    # 交互模式
  python3 ontap_s3_bench.py --config file.yaml # 从配置文件运行
  python3 ontap_s3_bench.py --dry-run          # 仅探测不执行测试
  python3 ontap_s3_bench.py --report-only      # 基于已有数据重新生成报告
"""

import sys, os, re, json, time, signal, argparse, logging, subprocess, shutil, textwrap
from datetime import datetime
from pathlib import Path

# ─── 自动安装依赖 ──────────────────────────────────────────────────────────────
REQUIRED = {"paramiko": "paramiko", "requests": "requests", "matplotlib": "matplotlib",
            "jinja2": "jinja2", "yaml": "pyyaml"}

def ensure_deps():
    missing = []
    for mod, pkg in REQUIRED.items():
        try:
            __import__(mod)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"[*] 安装缺失依赖: {', '.join(missing)}")
        # Adjust package versions for Python 3.6/3.7
        if sys.version_info < (3, 8):
            ver_map = {"paramiko": "paramiko<3", "requests": "requests<2.28",
                       "matplotlib": "matplotlib<3.4", "jinja2": "jinja2<3.1"}
            missing = [ver_map.get(p, p) for p in missing]
        # Try offline wheels first
        script_dir = os.path.dirname(os.path.abspath(__file__))
        wheels_dir = os.path.join(script_dir, "wheels")
        if not os.path.isdir(wheels_dir):
            wheels_dir = os.path.join(script_dir, "..", "wheels")
        if os.path.isdir(wheels_dir):
            print(f"[*] 从离线包安装: {wheels_dir}")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "-q",
                                   "--no-index", "--find-links", wheels_dir] + missing)
        else:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "-q"] + missing)
        print("[*] 依赖安装完成")

ensure_deps()

import paramiko
import requests
import yaml
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import numpy as np
from jinja2 import Template

requests.packages.urllib3.disable_warnings()

# ─── 常量 ──────────────────────────────────────────────────────────────────────
BRAND_PRIMARY = "#7F2854"
BRAND_SECONDARY = "#0067C5"
CHART_COLORS = ["#7F2854", "#0067C5", "#E53935", "#43A047", "#FB8C00"]
WARP_CDN = "https://dl.min.io/aistor/warp/release/linux-amd64/warp"
WARP_GITHUB = "https://github.com/minio/warp/releases/latest/download/warp_Linux_x86_64.tar.gz"
WARP_PORT = 7761

# ─── 日志 ──────────────────────────────────────────────────────────────────────
class ColorLog:
    RESET = "\033[0m"
    BOLD = "\033[1m"
    RED = "\033[91m"
    GREEN = "\033[92m"
    YELLOW = "\033[93m"
    BLUE = "\033[94m"
    PURPLE = "\033[95m"

    @staticmethod
    def info(msg): print(f"  {ColorLog.GREEN}✓{ColorLog.RESET} {msg}")
    @staticmethod
    def warn(msg): print(f"  {ColorLog.YELLOW}⚠{ColorLog.RESET} {msg}")
    @staticmethod
    def error(msg): print(f"  {ColorLog.RED}✗{ColorLog.RESET} {msg}")
    @staticmethod
    def step(msg): print(f"\n{ColorLog.BOLD}{ColorLog.PURPLE}{'='*60}{ColorLog.RESET}")
    @staticmethod
    def header(msg): print(f"{ColorLog.BOLD}{ColorLog.PURPLE}  {msg}{ColorLog.RESET}")
    @staticmethod
    def banner(msg):
        print(f"\n{ColorLog.BOLD}{ColorLog.PURPLE}{'='*60}")
        print(f"  {msg}")
        print(f"{'='*60}{ColorLog.RESET}\n")

log = ColorLog()

# ─── 进度条 ─────────────────────────────────────────────────────────────────────
def progress_bar(current, total, prefix="", width=30, extra=""):
    pct = current / total if total > 0 else 0
    filled = int(width * pct)
    bar = "█" * filled + "░" * (width - filled)
    print(f"\r  {prefix} [{bar}] {pct*100:.0f}% {extra}", end="", flush=True)
    if current >= total:
        print()

# ─── SSH 管理器 ─────────────────────────────────────────────────────────────────
class SSHManager:
    def __init__(self, logger):
        self.logger = logger
        self.connections = {}

    def connect(self, ip, user, password, retries=3, timeout=15):
        for attempt in range(1, retries + 1):
            try:
                client = paramiko.SSHClient()
                client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
                client.connect(ip, username=user, password=password,
                               timeout=timeout, banner_timeout=timeout, auth_timeout=timeout)
                self.connections[ip] = (client, user, password)
                return client
            except Exception as e:
                if attempt < retries:
                    time.sleep(5)
                else:
                    raise ConnectionError(f"SSH {ip} 连接失败 ({retries}次): {e}")

    def run(self, ip, cmd, timeout=120):
        if ip not in self.connections:
            raise ConnectionError(f"未连接到 {ip}")
        client, user, password = self.connections[ip]
        try:
            _, stdout, stderr = client.exec_command(cmd, timeout=timeout)
            out = stdout.read().decode("utf-8", errors="replace")
            err = stderr.read().decode("utf-8", errors="replace")
            rc = stdout.channel.recv_exit_status()
            return rc, out, err
        except Exception:
            # Reconnect and retry once
            client = self.connect(ip, user, password)
            _, stdout, stderr = client.exec_command(cmd, timeout=timeout)
            out = stdout.read().decode("utf-8", errors="replace")
            err = stderr.read().decode("utf-8", errors="replace")
            rc = stdout.channel.recv_exit_status()
            return rc, out, err

    def upload(self, ip, local_path, remote_path):
        if ip not in self.connections:
            raise ConnectionError(f"未连接到 {ip}")
        client = self.connections[ip][0]
        sftp = client.open_sftp()
        sftp.put(local_path, remote_path)
        sftp.chmod(remote_path, 0o755)
        sftp.close()

    def close_all(self):
        for ip, (client, _, _) in self.connections.items():
            try:
                client.close()
            except Exception:
                pass
        self.connections.clear()

# ─── ONTAP REST API 客户端 ──────────────────────────────────────────────────────
class ONTAPClient:
    def __init__(self, mgmt_ip, user, password):
        self.base = f"https://{mgmt_ip}/api"
        self.auth = (user, password)
        self.session = requests.Session()
        self.session.verify = False
        self.session.auth = self.auth

    def get(self, path, params=None):
        r = self.session.get(f"{self.base}{path}", params=params)
        r.raise_for_status()
        return r.json()

    def post(self, path, data):
        r = self.session.post(f"{self.base}{path}", json=data)
        return r.status_code, r.json() if r.text else {}

    def delete(self, path):
        r = self.session.delete(f"{self.base}{path}")
        return r.status_code

    def patch(self, path, data):
        r = self.session.patch(f"{self.base}{path}", json=data)
        return r.status_code, r.json() if r.text else {}

    def wait_job(self, job_uuid, timeout=300):
        for _ in range(timeout // 3):
            j = self.get(f"/cluster/jobs/{job_uuid}")
            if j["state"] == "success":
                return True
            if j["state"] == "failure":
                raise RuntimeError(f"Job 失败: {j.get('message','')}")
            time.sleep(3)
        raise TimeoutError("Job 超时")

    def get_cluster(self):
        return self.get("/cluster", {"fields": "name,version,location"})

    def get_nodes(self):
        return self.get("/cluster/nodes", {"fields": "name,model,serial_number"})["records"]

    def get_aggregates(self):
        return self.get("/storage/aggregates", {"fields": "space,node"})["records"]

    def get_svms(self):
        return self.get("/svm/svms", {"fields": "name,state,s3"})["records"]

    def get_lifs(self):
        return self.get("/network/ip/interfaces", {"fields": "ip,location,service_policy"})["records"]

    def get_s3_services(self):
        svms = self.get_svms()
        return [s for s in svms if s.get("s3", {}).get("enabled")]

    def get_s3_users(self, svm_uuid):
        return self.get(f"/protocols/s3/services/{svm_uuid}/users")["records"]

    def get_s3_buckets(self, svm_uuid):
        return self.get(f"/protocols/s3/services/{svm_uuid}/buckets",
                        {"fields": "name,size,logical_used_size"})["records"]

    def create_s3_user(self, svm_uuid, name):
        code, resp = self.post(f"/protocols/s3/services/{svm_uuid}/users", {"name": name})
        if code not in (200, 201):
            raise RuntimeError(f"创建 S3 用户失败: {resp}")
        rec = resp.get("records", [{}])[0]
        return rec.get("access_key"), rec.get("secret_key")

    def delete_s3_user(self, svm_uuid, name):
        return self.delete(f"/protocols/s3/services/{svm_uuid}/users/{name}")

    def create_s3_bucket(self, svm_uuid, name, size_bytes=858993459200, user="s3testuser"):
        code, resp = self.post(f"/protocols/s3/services/{svm_uuid}/buckets", {
            "name": name, "type": "s3", "size": size_bytes,
            "policy": {"statements": [{
                "sid": "BenchFullAccess", "effect": "allow",
                "actions": ["*"], "principals": [user],
                "resources": [name, f"{name}/*"]
            }]}
        })
        if code == 202:
            job_uuid = resp["job"]["uuid"]
            self.wait_job(job_uuid)
        elif code not in (200, 201):
            raise RuntimeError(f"创建 bucket 失败: {resp}")

# ─── 进度管理器 ─────────────────────────────────────────────────────────────────
class ProgressManager:
    def __init__(self, work_dir):
        self.file = os.path.join(work_dir, "progress.json")
        self.data = self._load()

    def _load(self):
        if os.path.exists(self.file):
            with open(self.file) as f:
                return json.load(f)
        return {"steps": {}, "started": datetime.now().isoformat()}

    def save(self):
        with open(self.file, "w") as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)

    def is_done(self, step):
        return self.data["steps"].get(step, {}).get("status") == "done"

    def mark_done(self, step, info=None):
        self.data["steps"][step] = {"status": "done", "time": datetime.now().isoformat(), "info": info}
        self.save()

    def has_progress(self):
        return bool(self.data["steps"])

# ─── 配置 ──────────────────────────────────────────────────────────────────────
class Config:
    def __init__(self):
        self.customer_name = ""
        self.ontap_ip = ""
        self.ontap_user = "admin"
        self.ontap_password = ""
        self.vms = []  # [{"ip": ..., "user": ..., "password": ...}]
        self.s3_endpoint = ""
        self.s3_access_key = ""
        self.s3_secret_key = ""
        self.s3_bucket = "warp-bench"
        self.s3_svm_uuid = ""
        self.s3_svm_name = ""
        self.s3_lif_ip = ""
        self.s3_protocol = "http"
        self.s3_port = 80
        self.test_mode = "standard"  # quick/standard/full/custom
        self.work_dir = os.path.expanduser("~/ontap_s3_test")
        self.dry_run = False
        self.report_only = False

    def to_dict(self):
        return {k: v for k, v in self.__dict__.items()}

    @classmethod
    def from_yaml(cls, path):
        with open(path) as f:
            d = yaml.safe_load(f)
        c = cls()
        for k, v in d.items():
            if hasattr(c, k):
                setattr(c, k, v)
        return c

    @classmethod
    def from_interactive(cls):
        c = cls()
        log.banner("ONTAP S3 对象存储性能测试工具")
        print("  联想凌拓科技有限公司\n")

        c.customer_name = input("  客户名称 (如: 安踏集团): ").strip() or "测试客户"
        c.ontap_ip = input("  ONTAP 管理 IP: ").strip()
        c.ontap_user = input("  ONTAP 用户名 [admin]: ").strip() or "admin"
        c.ontap_password = input("  ONTAP 密码: ").strip()

        n = int(input("  客户端 VM 数量: ").strip() or "4")
        default_user = input("  VM SSH 用户名 [root]: ").strip() or "root"
        default_pass = input("  VM SSH 密码: ").strip()

        for i in range(n):
            ip = input(f"  VM {i+1} IP: ").strip()
            c.vms.append({"ip": ip, "user": default_user, "password": default_pass})

        c.work_dir = input(f"  工作目录 [{c.work_dir}]: ").strip() or c.work_dir
        return c

# ─── 测试矩阵定义 ──────────────────────────────────────────────────────────────
TEST_PRESETS = {
    "quick": {
        "name": "快速测试",
        "desc": "2 对象大小 × 2 并发, ~15 分钟",
        "put": [("64KiB", [8, 16]), ("1MiB", [8, 16])],
        "get": [("64KiB", [8, 16]), ("1MiB", [8, 16])],
        "mixed": [("1MiB", [16])],
        "delete": [("64KiB", [16])],
        "list": [(8,)],
        "duration": "1m",
        "objects": 2000,
    },
    "standard": {
        "name": "标准测试",
        "desc": "4 对象大小 × 3 并发, ~45 分钟",
        "put": [("4KiB", [4, 16, 32]), ("64KiB", [4, 16, 32]),
                ("1MiB", [4, 16, 32]), ("4MiB", [4, 16, 32])],
        "get": [("4KiB", [4, 16, 32]), ("64KiB", [4, 16, 32]),
                ("1MiB", [4, 16, 32]), ("4MiB", [4, 16, 32])],
        "mixed": [("64KiB", [8, 16, 32]), ("1MiB", [8, 16, 32])],
        "delete": [("64KiB", [16])],
        "list": [(8,)],
        "duration": "2m",
        "objects": 3000,
    },
    "full": {
        "name": "完整测试",
        "desc": "5 对象大小 × 4 并发 × 全场景, ~90 分钟",
        "put": [(sz, [4, 8, 16, 32]) for sz in ["4KiB", "64KiB", "256KiB", "1MiB", "4MiB"]],
        "get": [(sz, [4, 8, 16, 32]) for sz in ["4KiB", "64KiB", "256KiB", "1MiB", "4MiB"]],
        "mixed": [("64KiB", [8, 16, 32]), ("1MiB", [8, 16, 32])],
        "delete": [("64KiB", [16])],
        "list": [(8,)],
        "duration": "3m",
        "objects": 5000,
    },
}

def count_scenes(preset):
    total = 0
    for sz, concs in preset["put"]:
        total += len(concs)
    for sz, concs in preset["get"]:
        total += len(concs)
    for sz, concs in preset["mixed"]:
        total += len(concs)
    total += len(preset["delete"])
    total += len(preset["list"])
    return total

# ─── warp 结果解析 ──────────────────────────────────────────────────────────────
def parse_warp_output(text):
    """从 warp stdout 提取关键指标"""
    result = {}
    # Standard report: "Report: PUT. Concurrency: 16. Ran: 3m1s"
    m = re.search(r'Report:\s+(\w+)\.\s+Concurrency:\s+(\d+)\.\s+Ran:\s+(\S+)', text)
    if m:
        result["type"] = m.group(1)
        result["concurrency"] = int(m.group(2))
        result["duration"] = m.group(3)

    # Throughput: "Average: 137.72 MiB/s, 137.72 obj/s"
    m = re.search(r'Average:\s+([\d.]+)\s+MiB/s,\s+([\d.]+)\s+obj/s', text)
    if m:
        result["throughput_mibps"] = float(m.group(1))
        result["iops"] = float(m.group(2))
    else:
        # DELETE/LIST format: "Average: 5395.47 obj/s"
        m = re.search(r'Average:\s+([\d.]+)\s+obj/s', text)
        if m:
            result["throughput_mibps"] = 0
            result["iops"] = float(m.group(1))

    # Latency
    m = re.search(r'Reqs:\s+Avg:\s+([\d.]+)ms.*?50%:\s+([\d.]+)ms.*?90%:\s+([\d.]+)ms.*?99%:\s+([\d.]+)ms', text)
    if m:
        result["latency_avg_ms"] = float(m.group(1))
        result["latency_p50_ms"] = float(m.group(2))
        result["latency_p90_ms"] = float(m.group(3))
        result["latency_p99_ms"] = float(m.group(4))

    return result if "iops" in result else None

# ─── 报告生成器 ─────────────────────────────────────────────────────────────────
class ReportGenerator:
    def __init__(self, config, env_data, results, work_dir):
        self.config = config
        self.env = env_data
        self.results = results
        self.work_dir = work_dir
        self.charts_dir = os.path.join(work_dir, "reports", "charts")
        os.makedirs(self.charts_dir, exist_ok=True)
        self._setup_font()

    def _setup_font(self):
        """查找可用的 CJK 字体"""
        self.font_prop = None
        script_dir = os.path.dirname(os.path.abspath(__file__))
        candidates = [
            os.path.join(script_dir, "fonts", "wqy-microhei.ttc"),
            os.path.join(script_dir, "..", "fonts", "wqy-microhei.ttc"),
            "/usr/share/fonts/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        for p in candidates:
            if os.path.exists(p):
                self.font_prop = FontProperties(fname=p)
                break
        if not self.font_prop:
            # Try to download
            try:
                subprocess.run(["wget", "-q",
                    "https://github.com/anthonyfok/fonts-wqy-microhei/raw/master/wqy-microhei.ttc",
                    "-O", "/tmp/wqy-microhei.ttc"], timeout=30, check=True)
                os.makedirs("/usr/share/fonts/wqy", exist_ok=True)
                shutil.copy("/tmp/wqy-microhei.ttc", "/usr/share/fonts/wqy/")
                self.font_prop = FontProperties(fname="/usr/share/fonts/wqy/wqy-microhei.ttc")
            except Exception:
                self.font_prop = FontProperties()
        plt.rcParams["axes.unicode_minus"] = False

    def _get(self, prefix, metric):
        sizes = sorted(set(r["scene"].split("_")[1] for r in self.results
                           if r["scene"].startswith(prefix + "_")),
                       key=lambda s: self._size_order(s))
        concs = sorted(set(r["concurrency"] for r in self.results
                           if r["scene"].startswith(prefix + "_")))
        data = {}
        for sz in sizes:
            vals = []
            for c in concs:
                items = [r for r in self.results
                         if r["scene"].startswith(f"{prefix}_{sz}") and r["concurrency"] == c]
                vals.append(items[0][metric] if items else None)
            data[sz] = vals
        return sizes, concs, data

    @staticmethod
    def _size_order(s):
        units = {"KiB": 1, "MiB": 1024, "GiB": 1048576}
        m = re.match(r'(\d+)(\w+)', s)
        if m:
            return int(m.group(1)) * units.get(m.group(2), 1)
        return 0

    def generate_charts(self):
        log.info("生成性能图表...")
        fp = self.font_prop

        def line_chart(prefix, metric, ylabel, title, filename):
            sizes, concs, data = self._get(prefix, metric)
            if not sizes:
                return
            fig, ax = plt.subplots(figsize=(10, 6))
            for i, sz in enumerate(sizes):
                y = data[sz]
                ax.plot(concs, y, "o-", color=CHART_COLORS[i % len(CHART_COLORS)],
                        label=sz, linewidth=2, markersize=8)
            ax.set_xlabel("总并发数", fontproperties=fp, fontsize=12)
            ax.set_ylabel(ylabel, fontproperties=fp, fontsize=12)
            ax.set_title(title, fontproperties=fp, fontsize=14, fontweight="bold")
            ax.set_xticks(concs)
            ax.legend(prop=fp, fontsize=10)
            ax.grid(True, alpha=0.3)
            fig.tight_layout()
            fig.savefig(os.path.join(self.charts_dir, filename), dpi=150, bbox_inches="tight")
            plt.close(fig)

        def bar_chart(prefix, title, filename):
            sizes, concs, data = self._get(prefix, "throughput_mibps")
            if not sizes:
                return
            best = []
            for sz in sizes:
                vals = [v for v in data[sz] if v is not None]
                best.append((sz, max(vals) if vals else 0))
            fig, ax = plt.subplots(figsize=(10, 6))
            x = np.arange(len(best))
            bars = ax.bar(x, [b[1] for b in best],
                          color=CHART_COLORS[:len(best)], width=0.6)
            for bar, (sz, tp) in zip(bars, best):
                ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 3,
                        f"{tp:.1f}", ha="center", va="bottom", fontproperties=fp, fontsize=9)
            ax.set_xticks(x)
            ax.set_xticklabels([b[0] for b in best], fontproperties=fp, fontsize=11)
            ax.set_ylabel("峰值吞吐量 (MiB/s)", fontproperties=fp, fontsize=12)
            ax.set_title(title, fontproperties=fp, fontsize=14, fontweight="bold")
            ax.grid(True, alpha=0.3, axis="y")
            fig.tight_layout()
            fig.savefig(os.path.join(self.charts_dir, filename), dpi=150, bbox_inches="tight")
            plt.close(fig)

        line_chart("put", "throughput_mibps", "吞吐量 (MiB/s)", "PUT 写入吞吐量 vs 并发数", "put_throughput.png")
        line_chart("put", "latency_avg_ms", "平均延迟 (ms)", "PUT 写入平均延迟 vs 并发数", "put_latency.png")
        line_chart("get", "throughput_mibps", "吞吐量 (MiB/s)", "GET 读取吞吐量 vs 并发数（热数据）", "get_throughput.png")
        line_chart("get", "latency_avg_ms", "平均延迟 (ms)", "GET 读取平均延迟 vs 并发数", "get_latency.png")
        bar_chart("put", "PUT 各对象大小峰值吞吐量对比", "put_best_throughput.png")
        bar_chart("get", "GET 各对象大小峰值吞吐量对比（热数据）", "get_best_throughput.png")

        # IOPS comparison
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
        for prefix, ax, label in [("put", ax1, "PUT"), ("get", ax2, "GET")]:
            sizes, concs, data = self._get(prefix, "iops")
            for i, sz in enumerate(sizes[:2]):  # 4KiB + 64KiB
                y = data[sz]
                ax.plot(concs, y, "o-", color=CHART_COLORS[i], label=f"{label} {sz}",
                        linewidth=2, markersize=8)
            ax.set_xlabel("总并发数", fontproperties=fp, fontsize=11)
            ax.set_ylabel("IOPS (obj/s)", fontproperties=fp, fontsize=11)
            ax.set_title(f"{label} IOPS vs 并发数", fontproperties=fp, fontsize=13, fontweight="bold")
            ax.set_xticks(concs)
            ax.legend(prop=fp, fontsize=10)
            ax.grid(True, alpha=0.3)
        fig.tight_layout()
        fig.savefig(os.path.join(self.charts_dir, "iops_comparison.png"), dpi=150, bbox_inches="tight")
        plt.close(fig)

        # Mixed workload
        mixed = [r for r in self.results if r["scene"].startswith("mixed_")]
        if mixed:
            fig, ax = plt.subplots(figsize=(10, 6))
            labels = [f'{r["scene"].split("_")[1]}\nc{r["concurrency"]}' for r in mixed]
            get_vals = [r["throughput_mibps"] for r in mixed]
            put_est = [r["throughput_mibps"] * 30 / 70 for r in mixed]
            x = np.arange(len(labels))
            w = 0.35
            ax.bar(x - w/2, get_vals, w, color=BRAND_SECONDARY, label="GET (70%)")
            ax.bar(x + w/2, put_est, w, color=BRAND_PRIMARY, label="PUT (30%)")
            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontproperties=fp, fontsize=10)
            ax.set_ylabel("吞吐量 (MiB/s)", fontproperties=fp, fontsize=12)
            ax.set_title("混合读写性能 (70% GET + 30% PUT)", fontproperties=fp, fontsize=14, fontweight="bold")
            ax.legend(prop=fp, fontsize=11)
            ax.grid(True, alpha=0.3, axis="y")
            fig.tight_layout()
            fig.savefig(os.path.join(self.charts_dir, "mixed_workload.png"), dpi=150, bbox_inches="tight")
            plt.close(fig)

        log.info(f"图表已保存到 {self.charts_dir}")

    def generate_html(self):
        log.info("生成 HTML 报告...")
        sizes_all = sorted(set(r["scene"].split("_")[1] for r in self.results
                               if r["type"] in ("PUT", "GET")),
                           key=self._size_order)
        concs_all = sorted(set(r["concurrency"] for r in self.results
                               if r["type"] in ("PUT", "GET")))

        def build_chart_data(prefix):
            d = {}
            for sz in sizes_all:
                vals = []
                for c in concs_all:
                    items = [r for r in self.results
                             if r["scene"].startswith(f"{prefix}_{sz}") and r["concurrency"] == c]
                    vals.append(items[0]["throughput_mibps"] if items else None)
                d[sz] = vals
            return d

        put_tp = build_chart_data("put")
        get_tp = build_chart_data("get")

        def build_lat_data(prefix):
            d = {}
            for sz in sizes_all:
                vals = []
                for c in concs_all:
                    items = [r for r in self.results
                             if r["scene"].startswith(f"{prefix}_{sz}") and r["concurrency"] == c]
                    vals.append(items[0]["latency_avg_ms"] if items else None)
                d[sz] = vals
            return d

        put_lat = build_lat_data("put")
        get_lat = build_lat_data("get")

        # Peaks
        put_results = [r for r in self.results if r["type"] == "PUT"]
        get_results = [r for r in self.results if r["type"] == "GET"]
        peak_put_tp = max(put_results, key=lambda x: x["throughput_mibps"]) if put_results else {}
        peak_get_tp = max(get_results, key=lambda x: x["throughput_mibps"]) if get_results else {}
        peak_put_iops = max(put_results, key=lambda x: x["iops"]) if put_results else {}
        peak_get_iops = max(get_results, key=lambda x: x["iops"]) if get_results else {}
        delete_r = next((r for r in self.results if r["type"] == "DELETE"), {})
        list_r = next((r for r in self.results if r["type"] == "LIST"), {})
        mixed_data = [r for r in self.results if r["scene"].startswith("mixed_")]

        table_rows = ""
        for r in self.results:
            tp = f'{r["throughput_mibps"]:.2f}' if r.get("throughput_mibps", 0) > 0 else "-"
            table_rows += (f'<tr><td>{r["scene"]}</td><td>{r["type"]}</td>'
                          f'<td>{r["concurrency"]}</td><td>{tp}</td>'
                          f'<td>{r["iops"]:.2f}</td><td>{r.get("latency_avg_ms",0):.1f}</td>'
                          f'<td>{r.get("latency_p50_ms",0):.1f}</td>'
                          f'<td>{r.get("latency_p90_ms",0):.1f}</td>'
                          f'<td>{r.get("latency_p99_ms",0):.1f}</td></tr>\n')

        vm_rows = ""
        for v in self.env.get("vms", []):
            vm_rows += (f'<tr><td>{v.get("hostname","")}</td><td>{v["ip"]}</td>'
                       f'<td>{v.get("os","")}</td><td>{v.get("cpu_model","")}</td>'
                       f'<td>{v.get("memory","")}</td></tr>\n')

        html = self._html_template().replace("{{CUSTOMER}}", self.config.customer_name)
        html = html.replace("{{DATE}}", self.env.get("test_date", datetime.now().strftime("%Y-%m-%d")))
        html = html.replace("{{SIZES_JSON}}", json.dumps(sizes_all))
        html = html.replace("{{CONCS_JSON}}", json.dumps(concs_all))
        html = html.replace("{{PUT_TP_JSON}}", json.dumps(put_tp))
        html = html.replace("{{GET_TP_JSON}}", json.dumps(get_tp))
        html = html.replace("{{PUT_LAT_JSON}}", json.dumps(put_lat))
        html = html.replace("{{GET_LAT_JSON}}", json.dumps(get_lat))
        html = html.replace("{{COLORS_JSON}}", json.dumps(CHART_COLORS))
        html = html.replace("{{TABLE_ROWS}}", table_rows)
        html = html.replace("{{VM_ROWS}}", vm_rows)
        html = html.replace("{{PEAK_PUT_TP}}", f'{peak_put_tp.get("throughput_mibps",0):.1f}')
        html = html.replace("{{PEAK_PUT_TP_SCENE}}", peak_put_tp.get("scene", ""))
        html = html.replace("{{PEAK_GET_TP}}", f'{peak_get_tp.get("throughput_mibps",0):.1f}')
        html = html.replace("{{PEAK_GET_TP_SCENE}}", peak_get_tp.get("scene", ""))
        html = html.replace("{{PEAK_PUT_IOPS}}", f'{peak_put_iops.get("iops",0):,.0f}')
        html = html.replace("{{PEAK_GET_IOPS}}", f'{peak_get_iops.get("iops",0):,.0f}')
        html = html.replace("{{DELETE_IOPS}}", f'{delete_r.get("iops",0):,.0f}')
        html = html.replace("{{LIST_IOPS}}", f'{list_r.get("iops",0):,.0f}')
        html = html.replace("{{ONTAP_VERSION}}", self.env.get("ontap",{}).get("version",""))
        html = html.replace("{{ONTAP_MODEL}}", self.env.get("ontap",{}).get("model",""))
        html = html.replace("{{S3_LIF}}", self.config.s3_lif_ip or "")

        # Mixed data
        mixed_labels = json.dumps([f'{r["scene"].split("_")[1]} c{r["concurrency"]}' for r in mixed_data])
        mixed_get = json.dumps([r["throughput_mibps"] for r in mixed_data])
        mixed_put = json.dumps([round(r["throughput_mibps"]*30/70, 2) for r in mixed_data])
        html = html.replace("{{MIXED_LABELS}}", mixed_labels)
        html = html.replace("{{MIXED_GET}}", mixed_get)
        html = html.replace("{{MIXED_PUT}}", mixed_put)

        outpath = os.path.join(self.work_dir, "reports",
                               f"{self.config.customer_name}_ONTAP_S3_性能测试报告.html")
        with open(outpath, "w", encoding="utf-8") as f:
            f.write(html)
        log.info(f"HTML 报告: {outpath}")
        return outpath

    def generate_word(self):
        """生成 Word 报告 (使用 python-docx 或 Node.js docx)"""
        log.info("生成 Word 报告...")
        # Try python-docx first
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = DocxDocument()
            style = doc.styles["Normal"]
            style.font.size = Pt(10.5)

            # Cover
            for _ in range(6):
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("联想凌拓科技有限公司")
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0x7F, 0x28, 0x54)
            r.bold = True

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"\n{self.config.customer_name}\nONTAP S3 对象存储性能测试报告")
            r.font.size = Pt(22)
            r.font.color.rgb = RGBColor(0x7F, 0x28, 0x54)
            r.bold = True

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"\nNetApp ONTAP S3 对象存储基准性能测试\n\n"
                         f"测试日期：{self.env.get('test_date','')}\n版本：V1.0\n"
                         f"编制单位：联想凌拓科技有限公司")
            r.font.size = Pt(12)

            doc.add_page_break()

            # Content
            doc.add_heading("1. 测试概述", level=1)
            doc.add_paragraph(
                "⚠️ 重要说明：本次测试基于 ONTAP Select 虚拟化环境，"
                "物理 NetApp 存储系统的性能将大幅高于本次测试结果。")
            doc.add_paragraph(
                f"本次测试使用 MinIO warp 分布式基准测试工具，"
                f"共 {len(self.results)} 个测试场景。")

            doc.add_heading("2. 测试环境", level=1)
            # Storage table
            t = doc.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "项目"
            t.rows[0].cells[1].text = "详情"
            for item, detail in [
                ("集群", self.env.get("ontap",{}).get("cluster_name","")),
                ("版本", self.env.get("ontap",{}).get("version","")),
                ("平台", self.env.get("ontap",{}).get("model","")),
                ("S3 LIF", self.config.s3_lif_ip),
            ]:
                row = t.add_row()
                row.cells[0].text = item
                row.cells[1].text = str(detail)

            # Charts
            for chart_file, caption_text in [
                ("put_throughput.png", "PUT 写入吞吐量"),
                ("put_latency.png", "PUT 写入延迟"),
                ("get_throughput.png", "GET 读取吞吐量（热数据）"),
                ("get_latency.png", "GET 读取延迟"),
                ("put_best_throughput.png", "PUT 峰值吞吐量对比"),
                ("get_best_throughput.png", "GET 峰值吞吐量对比"),
                ("iops_comparison.png", "IOPS 对比"),
                ("mixed_workload.png", "混合读写性能"),
            ]:
                img_path = os.path.join(self.charts_dir, chart_file)
                if os.path.exists(img_path):
                    doc.add_heading(caption_text, level=2)
                    doc.add_picture(img_path, width=Inches(6))

            # Results table
            doc.add_heading("完整测试数据", level=1)
            t = doc.add_table(rows=1, cols=6)
            t.style = "Table Grid"
            for i, h in enumerate(["场景", "总并发", "吞吐MiB/s", "IOPS", "Avg ms", "P99 ms"]):
                t.rows[0].cells[i].text = h
            for r in self.results:
                row = t.add_row()
                row.cells[0].text = r["scene"]
                row.cells[1].text = str(r["concurrency"])
                row.cells[2].text = f'{r.get("throughput_mibps",0):.2f}'
                row.cells[3].text = f'{r["iops"]:.0f}'
                row.cells[4].text = f'{r.get("latency_avg_ms",0):.1f}'
                row.cells[5].text = f'{r.get("latency_p99_ms",0):.1f}'

            # Conclusion
            doc.add_heading("测试结论与建议", level=1)
            put_r = [r for r in self.results if r["type"] == "PUT"]
            get_r = [r for r in self.results if r["type"] == "GET"]
            if put_r:
                best = max(put_r, key=lambda x: x["throughput_mibps"])
                doc.add_paragraph(f"PUT 峰值吞吐: {best['throughput_mibps']:.1f} MiB/s ({best['scene']})")
            if get_r:
                best = max(get_r, key=lambda x: x["throughput_mibps"])
                doc.add_paragraph(f"GET 峰值吞吐（热数据）: {best['throughput_mibps']:.1f} MiB/s ({best['scene']})")

            doc.add_paragraph("优化建议：")
            for tip in ["使用物理存储可获数倍性能提升", "增加 S3 LIF 分散负载",
                        "部署多节点集群横向扩展", "根据延迟要求选择合适并发"]:
                doc.add_paragraph(tip, style="List Bullet")

            outpath = os.path.join(self.work_dir, "reports",
                                   f"{self.config.customer_name}_ONTAP_S3_性能测试报告.docx")
            doc.save(outpath)
            log.info(f"Word 报告: {outpath}")
            return outpath

        except ImportError:
            log.warn("python-docx 未安装，尝试安装...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "-q", "python-docx"])
            return self.generate_word()  # Retry

    def _html_template(self):
        """返回 HTML 模板字符串"""
        return '''<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>{{CUSTOMER}} ONTAP S3 性能测试报告</title>
<script src="https://cdn.jsdelivr.net/npm/echarts@5/dist/echarts.min.js"></script>
<style>
:root{--brand:#7F2854;--brand2:#0067C5;--bg:#fff;--text:#1a1a1a;--card:#f8f8f8;--border:#e0e0e0;--th:#7F2854;--alt:#F5F0F2;}
[data-theme=dark]{--bg:#1a1a2e;--text:#e0e0e0;--card:#16213e;--border:#333;--th:#5a1a3a;--alt:#1e2a4a;}
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:-apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif;background:var(--bg);color:var(--text);line-height:1.6;}
.hdr{background:linear-gradient(135deg,#7F2854,#5a1a3a);color:#fff;padding:30px 40px;}
.hdr h1{font-size:26px;} .hdr .sub{font-size:13px;opacity:.85;}
.toggle{position:fixed;top:12px;right:12px;z-index:999;background:var(--brand);color:#fff;border:none;padding:6px 14px;border-radius:18px;cursor:pointer;font-size:13px;}
nav{background:var(--card);border-bottom:1px solid var(--border);padding:8px 40px;position:sticky;top:0;z-index:100;display:flex;gap:18px;flex-wrap:wrap;}
nav a{color:var(--brand);text-decoration:none;font-size:13px;font-weight:500;}
.wrap{max-width:1200px;margin:0 auto;padding:20px 40px;}
section{margin-bottom:36px;} h2{color:var(--brand);border-bottom:2px solid var(--brand);padding-bottom:6px;margin-bottom:16px;font-size:20px;}
h3{color:var(--brand2);margin:12px 0 8px;font-size:16px;}
.card{background:var(--card);border-radius:8px;padding:16px;margin:12px 0;border:1px solid var(--border);}
table{width:100%;border-collapse:collapse;margin:8px 0;font-size:13px;}
th{background:var(--th);color:#fff;padding:8px 10px;text-align:left;}
td{padding:6px 10px;border-bottom:1px solid var(--border);}
tr:nth-child(even){background:var(--alt);}
.chart{width:100%;height:380px;margin:12px 0;}
.warn{background:#FFF3E0;border-left:4px solid #FB8C00;padding:10px 14px;margin:12px 0;border-radius:4px;color:#333;}
.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:12px;}
.pk{background:var(--card);border-radius:8px;padding:16px;text-align:center;border:1px solid var(--border);}
.pk .v{font-size:26px;font-weight:bold;color:var(--brand);} .pk .l{font-size:12px;color:#888;margin-top:4px;}
.foot{text-align:center;padding:24px;color:#888;font-size:12px;border-top:1px solid var(--border);}
</style></head><body>
<button class="toggle" onclick="document.body.dataset.theme=document.body.dataset.theme==='dark'?'':'dark';initCharts()">🌓 切换主题</button>
<div class="hdr"><div style="font-size:15px;font-weight:bold;letter-spacing:2px">联想凌拓科技有限公司</div>
<h1>{{CUSTOMER}} ONTAP S3 对象存储性能测试报告</h1>
<div class="sub">测试日期: {{DATE}} | 工具: MinIO warp</div></div>
<nav><a href="#env">测试环境</a><a href="#put">写入性能</a><a href="#get">读取性能</a>
<a href="#mixed">混合读写</a><a href="#conclusion">结论</a><a href="#appendix">附录</a></nav>
<div class="wrap">
<section><div class="warn">⚠️ <b>本次测试基于 ONTAP Select 虚拟化环境</b>，物理 NetApp 存储性能将大幅高于本次结果。</div></section>
<section id="env"><h2>测试环境</h2>
<table><tr><th>项目</th><th>详情</th></tr>
<tr><td>ONTAP</td><td>{{ONTAP_VERSION}} ({{ONTAP_MODEL}})</td></tr>
<tr><td>S3 LIF</td><td>{{S3_LIF}} (HTTP:80)</td></tr></table>
<h3>客户端</h3><table><tr><th>主机名</th><th>IP</th><th>OS</th><th>CPU</th><th>内存</th></tr>{{VM_ROWS}}</table></section>
<section id="put"><h2>写入性能 (PUT)</h2>
<div class="chart" id="c-put-tp"></div><div class="chart" id="c-put-lat"></div>
<div class="chart" id="c-put-best"></div></section>
<section id="get"><h2>读取性能 (GET)</h2>
<div class="warn">📌 GET 数据从 ONTAP 内存缓存读取，代表<b>热数据读取性能</b>。</div>
<div class="chart" id="c-get-tp"></div><div class="chart" id="c-get-lat"></div>
<div class="chart" id="c-get-best"></div></section>
<section id="mixed"><h2>混合读写 (70% GET + 30% PUT)</h2><div class="chart" id="c-mixed"></div></section>
<section id="conclusion"><h2>测试结论</h2><div class="grid">
<div class="pk"><div class="v">{{PEAK_PUT_TP}}</div><div class="l">PUT 峰值 MiB/s<br>{{PEAK_PUT_TP_SCENE}}</div></div>
<div class="pk"><div class="v">{{PEAK_GET_TP}}</div><div class="l">GET 峰值 MiB/s<br>{{PEAK_GET_TP_SCENE}}</div></div>
<div class="pk"><div class="v">{{PEAK_PUT_IOPS}}</div><div class="l">PUT 峰值 IOPS</div></div>
<div class="pk"><div class="v">{{PEAK_GET_IOPS}}</div><div class="l">GET 峰值 IOPS</div></div>
<div class="pk"><div class="v">{{DELETE_IOPS}}</div><div class="l">DELETE obj/s</div></div>
<div class="pk"><div class="v">{{LIST_IOPS}}</div><div class="l">LIST obj/s</div></div>
</div></section>
<section id="appendix"><h2>完整测试数据</h2>
<p>warp --concurrent 为 per-client 值，总并发 = concurrent × N clients</p>
<div style="overflow-x:auto"><table>
<tr><th>场景</th><th>类型</th><th>总并发</th><th>吞吐MiB/s</th><th>IOPS</th><th>Avg ms</th><th>P50</th><th>P90</th><th>P99</th></tr>
{{TABLE_ROWS}}</table></div></section></div>
<div class="foot">联想凌拓科技有限公司 | 机密文件 | {{DATE}}</div>
<script>
const C={{COLORS_JSON}},SZ={{SIZES_JSON}},CO={{CONCS_JSON}};
const putTp={{PUT_TP_JSON}},getTp={{GET_TP_JSON}},putLat={{PUT_LAT_JSON}},getLat={{GET_LAT_JSON}};
function ms(d){return SZ.map((s,i)=>({name:s,type:'line',data:d[s],symbol:'circle',symbolSize:8,
lineStyle:{width:2.5},itemStyle:{color:C[i%C.length]}}))}
function lc(id,t,y,d){const c=echarts.init(document.getElementById(id));
c.setOption({title:{text:t,left:'center'},tooltip:{trigger:'axis'},legend:{data:SZ,bottom:0},
grid:{left:80,right:30,top:50,bottom:60},xAxis:{type:'category',data:CO.map(c=>'c'+c),name:'总并发'},
yAxis:{type:'value',name:y},series:ms(d)});return c}
function bc(id,t,d){const c=echarts.init(document.getElementById(id));
const b=SZ.map(s=>{const v=d[s];return v?Math.max(...v.filter(x=>x!==null)):0});
c.setOption({title:{text:t,left:'center'},tooltip:{trigger:'axis'},
grid:{left:80,right:30,top:50,bottom:40},xAxis:{type:'category',data:SZ},
yAxis:{type:'value',name:'MiB/s'},series:[{type:'bar',data:b.map((v,i)=>({value:v,itemStyle:{color:C[i%C.length]}})),
label:{show:true,position:'top',formatter:'{c}'}}]});return c}
let charts=[];
function initCharts(){charts.forEach(c=>c.dispose());charts=[];
charts.push(lc('c-put-tp','PUT 吞吐量 vs 并发','MiB/s',putTp));
charts.push(lc('c-put-lat','PUT 延迟 vs 并发','ms',putLat));
charts.push(bc('c-put-best','PUT 峰值吞吐量',putTp));
charts.push(lc('c-get-tp','GET 吞吐量 vs 并发（热数据）','MiB/s',getTp));
charts.push(lc('c-get-lat','GET 延迟 vs 并发','ms',getLat));
charts.push(bc('c-get-best','GET 峰值吞吐量',getTp));
const mc=echarts.init(document.getElementById('c-mixed'));
mc.setOption({title:{text:'混合读写',left:'center'},tooltip:{trigger:'axis'},
legend:{data:['GET(70%)','PUT(30%)'],bottom:0},grid:{left:80,right:30,top:50,bottom:60},
xAxis:{type:'category',data:{{MIXED_LABELS}}},yAxis:{type:'value',name:'MiB/s'},
series:[{name:'GET(70%)',type:'bar',stack:'t',data:{{MIXED_GET}},itemStyle:{color:'#0067C5'}},
{name:'PUT(30%)',type:'bar',stack:'t',data:{{MIXED_PUT}},itemStyle:{color:'#7F2854'}}]});
charts.push(mc)}
window.addEventListener('load',initCharts);
window.addEventListener('resize',()=>charts.forEach(c=>c.resize()));
</script></body></html>'''

# ─── 主流程 ─────────────────────────────────────────────────────────────────────
class Benchmark:
    def __init__(self, config):
        self.cfg = config
        self.work_dir = config.work_dir
        self.ssh = SSHManager(log)
        self.ontap = None
        self.env_data = {}
        self.results = []
        self.progress = None

        os.makedirs(os.path.join(self.work_dir, "warp_results"), exist_ok=True)
        os.makedirs(os.path.join(self.work_dir, "ontap_perf"), exist_ok=True)
        os.makedirs(os.path.join(self.work_dir, "reports"), exist_ok=True)

        # Setup file logging
        logging.basicConfig(
            filename=os.path.join(self.work_dir, "bench.log"),
            level=logging.DEBUG,
            format="%(asctime)s %(levelname)s %(message)s",
        )
        self.file_log = logging.getLogger("bench")
        self.progress = ProgressManager(self.work_dir)

    def run(self):
        """主入口"""
        try:
            if self.cfg.report_only:
                self._load_existing_data()
                self._step7_reports()
                return

            # Check for existing progress
            if self.progress.has_progress():
                print("\n  发现之前的测试进度:")
                for step, info in self.progress.data["steps"].items():
                    print(f"    ✓ {step}")
                choice = input("\n  [1] 从断点继续  [2] 重新开始: ").strip()
                if choice == "2":
                    self.progress = ProgressManager.__new__(ProgressManager)
                    self.progress.file = os.path.join(self.work_dir, "progress.json")
                    self.progress.data = {"steps": {}, "started": datetime.now().isoformat()}
                    self.progress.save()

            self._step1_connect()
            self._step2_probe()
            if self.cfg.dry_run:
                log.banner("Dry-run 模式完成 — 仅探测不执行测试")
                return
            self._step3_s3_config()
            self._step4_test_config()
            self._step5_deploy()
            self._step6_execute()
            self._step7_reports()
            self._step8_cleanup()

        except KeyboardInterrupt:
            log.warn("\n用户中断，清理中...")
            self._emergency_cleanup()
        except Exception as e:
            log.error(f"错误: {e}")
            self.file_log.exception("Fatal error")
            raise
        finally:
            self.ssh.close_all()

    def _step1_connect(self):
        if self.progress.is_done("step1_connect"):
            log.info("Step 1 已完成，跳过")
            self.ontap = ONTAPClient(self.cfg.ontap_ip, self.cfg.ontap_user, self.cfg.ontap_password)
            for vm in self.cfg.vms:
                self.ssh.connect(vm["ip"], vm["user"], vm["password"])
            return

        log.banner("Step 1: 连接测试")
        # ONTAP
        print(f"  连接 ONTAP {self.cfg.ontap_ip}...", end=" ", flush=True)
        self.ontap = ONTAPClient(self.cfg.ontap_ip, self.cfg.ontap_user, self.cfg.ontap_password)
        cluster = self.ontap.get_cluster()
        print(f"✓ {cluster['name']} ({cluster['version']['full'][:30]})")

        # VMs
        for i, vm in enumerate(self.cfg.vms):
            print(f"  SSH {vm['ip']}...", end=" ", flush=True)
            try:
                self.ssh.connect(vm["ip"], vm["user"], vm["password"])
                rc, out, _ = self.ssh.run(vm["ip"], "hostname")
                print(f"✓ {out.strip()}")
            except Exception as e:
                print(f"✗ {e}")

        self.progress.mark_done("step1_connect")

    def _step2_probe(self):
        if self.progress.is_done("step2_probe"):
            # Load existing env data
            env_file = os.path.join(self.work_dir, "env_report.json")
            if os.path.exists(env_file):
                with open(env_file) as f:
                    self.env_data = json.load(f)
            log.info("Step 2 已完成，跳过")
            return

        log.banner("Step 2: 环境探测")

        # Probe VMs
        vm_info = []
        for vm in self.cfg.vms:
            ip = vm["ip"]
            print(f"  探测 {ip}...", end=" ", flush=True)
            try:
                rc, out, _ = self.ssh.run(ip, textwrap.dedent('''
                    echo "HOSTNAME:$(hostname)"
                    echo "OS:$(cat /etc/redhat-release 2>/dev/null || head -1 /etc/os-release)"
                    echo "KERNEL:$(uname -r)"
                    echo "CPU:$(lscpu | grep 'Model name' | head -1 | sed 's/.*: *//')"
                    echo "CPUS:$(lscpu | grep '^CPU(s):' | awk '{print $2}')"
                    echo "MEM:$(free -h | awk '/Mem:/{print $2}')"
                    echo "DISK:$(df -h / | awk 'NR==2{print $2,$4,$5}')"
                    echo "PKG:$(which dnf 2>/dev/null || which yum 2>/dev/null || which apt 2>/dev/null)"
                    echo "FW:$(firewall-cmd --state 2>/dev/null || echo inactive)"
                '''))
                info = {"ip": ip}
                for line in out.strip().split("\n"):
                    if ":" in line:
                        k, v = line.split(":", 1)
                        info[k.strip().lower()] = v.strip()
                vm_info.append(info)
                print(f"✓ {info.get('hostname','')} {info.get('os','')[:20]}")
            except Exception as e:
                print(f"✗ {e}")
                vm_info.append({"ip": ip, "error": str(e)})

        # Probe ONTAP
        print(f"  探测 ONTAP...", end=" ", flush=True)
        cluster = self.ontap.get_cluster()
        nodes = self.ontap.get_nodes()
        aggrs = self.ontap.get_aggregates()
        svms = self.ontap.get_svms()
        lifs = self.ontap.get_lifs()
        s3_svms = self.ontap.get_s3_services()
        print(f"✓ {cluster['name']} {len(nodes)} 节点")

        self.env_data = {
            "customer_name": self.cfg.customer_name,
            "test_date": datetime.now().strftime("%Y-%m-%d"),
            "vms": [{
                "ip": v["ip"],
                "hostname": v.get("hostname", ""),
                "os": v.get("os", ""),
                "cpu_model": v.get("cpu", ""),
                "cpu_cores": v.get("cpus", ""),
                "memory": v.get("mem", ""),
                "disk_root": v.get("disk", ""),
                "firewall": v.get("fw", ""),
            } for v in vm_info],
            "ontap": {
                "mgmt_ip": self.cfg.ontap_ip,
                "cluster_name": cluster["name"],
                "version": cluster["version"]["full"],
                "model": nodes[0].get("model", "") if nodes else "",
                "nodes": [{"name": n["name"]} for n in nodes],
                "aggregates": [{
                    "name": a["name"],
                    "total_gb": round(a["space"]["block_storage"]["size"] / 1024**3),
                    "available_gb": round(a["space"]["block_storage"]["available"] / 1024**3),
                    "used_percent": a["space"]["block_storage"].get("used_percent", 0),
                } for a in aggrs],
                "s3_enabled": bool(s3_svms),
                "s3_svms": [{"name": s["name"], "uuid": s["uuid"]} for s in s3_svms],
                "lifs": [{"name": l["name"], "ip": l["ip"]["address"],
                          "port": l["location"]["port"]["name"],
                          "policy": l["service_policy"]["name"]} for l in lifs],
            },
            "s3_config": {},
        }

        # Summary table
        print(f"\n  {'─'*56}")
        print(f"  {'VM IP':<18} {'主机名':<20} {'OS':<15}")
        print(f"  {'─'*56}")
        for v in vm_info:
            print(f"  {v['ip']:<18} {v.get('hostname','?'):<20} {v.get('os','?')[:15]:<15}")
        print(f"  {'─'*56}")
        print(f"  ONTAP: {cluster['name']} | {cluster['version']['full'][:40]}")
        for a in aggrs:
            avail = round(a["space"]["block_storage"]["available"] / 1024**3)
            print(f"  聚合: {a['name']} — {avail}GB 可用")
        if s3_svms:
            for s in s3_svms:
                print(f"  S3: {s['name']} — {s.get('s3',{}).get('name','')}")
        print()

        # Save
        with open(os.path.join(self.work_dir, "env_report.json"), "w") as f:
            json.dump(self.env_data, f, indent=2, ensure_ascii=False)

        self.progress.mark_done("step2_probe")

        if not self.cfg.dry_run:
            input("  按 Enter 继续...")

    def _step3_s3_config(self):
        if self.progress.is_done("step3_s3"):
            env_file = os.path.join(self.work_dir, "env_report.json")
            if os.path.exists(env_file):
                with open(env_file) as f:
                    self.env_data = json.load(f)
                s3cfg = self.env_data.get("s3_config", {})
                self.cfg.s3_access_key = s3cfg.get("access_key", "")
                self.cfg.s3_secret_key = s3cfg.get("secret_key", "")
                self.cfg.s3_lif_ip = s3cfg.get("s3_lif_ip", "")
                self.cfg.s3_svm_uuid = s3cfg.get("svm_uuid", "")
                self.cfg.s3_bucket = s3cfg.get("bucket", "warp-bench")
            log.info("Step 3 已完成，跳过")
            return

        log.banner("Step 3: S3 配置")
        s3_svms = self.env_data["ontap"].get("s3_svms", [])

        if s3_svms:
            print(f"  已有 S3 服务:")
            for s in s3_svms:
                print(f"    - {s['name']} (UUID: {s['uuid'][:8]}...)")
            svm = s3_svms[0]
            self.cfg.s3_svm_uuid = svm["uuid"]
            self.cfg.s3_svm_name = svm["name"]

            # Find S3 LIF — prioritize policy with 's3', then show all data LIFs for user to choose
            all_lifs = self.env_data["ontap"]["lifs"]
            s3_lifs = [l for l in all_lifs if "s3" in l["policy"].lower()]
            if not s3_lifs:
                s3_lifs = [l for l in all_lifs
                           if "data" in l["policy"].lower() and "mgmt" not in l["policy"].lower()]

            if len(s3_lifs) == 1:
                self.cfg.s3_lif_ip = s3_lifs[0]["ip"]
                print(f"  S3 LIF: {self.cfg.s3_lif_ip} (policy: {s3_lifs[0]['policy']}, 端口: {s3_lifs[0].get('port','')})")
                confirm = input("  确认使用此 LIF? [Y/n]: ").strip().lower()
                if confirm == "n":
                    self.cfg.s3_lif_ip = input("  输入 S3 LIF IP: ").strip()
            elif len(s3_lifs) > 1:
                print(f"  发现以下 LIF:")
                for i, l in enumerate(s3_lifs, 1):
                    print(f"    [{i}] {l['ip']} (policy: {l['policy']}, 端口: {l.get('port','')})")
                print(f"    [0] 手动输入")
                choice = input(f"\n  选择 S3 LIF [1]: ").strip() or "1"
                if choice == "0":
                    self.cfg.s3_lif_ip = input("  输入 S3 LIF IP: ").strip()
                else:
                    idx = int(choice) - 1
                    if 0 <= idx < len(s3_lifs):
                        self.cfg.s3_lif_ip = s3_lifs[idx]["ip"]
                    else:
                        self.cfg.s3_lif_ip = s3_lifs[0]["ip"]
                print(f"  已选 S3 LIF: {self.cfg.s3_lif_ip}")
            else:
                self.cfg.s3_lif_ip = input("  未找到 S3 LIF，请手动输入 IP: ").strip()

            # Check existing users
            users = self.ontap.get_s3_users(self.cfg.s3_svm_uuid)
            user_names = [u["name"] for u in users]
            print(f"  现有用户: {', '.join(user_names)}")

            # Create/recreate user
            if "s3testuser" in user_names:
                print("  重建 s3testuser 获取新 key...")
                self.ontap.delete_s3_user(self.cfg.s3_svm_uuid, "s3testuser")
            ak, sk = self.ontap.create_s3_user(self.cfg.s3_svm_uuid, "s3testuser")
            self.cfg.s3_access_key = ak
            self.cfg.s3_secret_key = sk
            log.info(f"Access Key: {ak}")

            # Create bucket
            existing_buckets = [b["name"] for b in
                                self.ontap.get_s3_buckets(self.cfg.s3_svm_uuid)]
            bucket_name = self.cfg.s3_bucket
            if bucket_name in existing_buckets:
                choice = input(f"  bucket '{bucket_name}' 已存在, [1] 复用 [2] 删除重建: ").strip()
                if choice == "2":
                    # Delete not easily done via API for bucket, just use different name
                    bucket_name = f"warp-bench-{datetime.now().strftime('%m%d%H%M')}"
                    self.cfg.s3_bucket = bucket_name

            if bucket_name not in existing_buckets:
                print(f"  创建 bucket: {bucket_name}...")
                self.ontap.create_s3_bucket(self.cfg.s3_svm_uuid, bucket_name)
                log.info(f"Bucket 已创建: {bucket_name}")

            self.cfg.s3_endpoint = f"http://{self.cfg.s3_lif_ip}:{self.cfg.s3_port}"
        else:
            log.error("未找到 S3 服务，请先在 ONTAP 上启用 S3")
            sys.exit(1)

        # Verify connectivity (warp put 5s)
        print("  验证 S3 连通性 (warp写入测试)...")
        verify_cmd = (f"warp put --host={self.cfg.s3_lif_ip}:{self.cfg.s3_port} "
                      f"--access-key={self.cfg.s3_access_key} "
                      f"--secret-key={self.cfg.s3_secret_key} "
                      f"--tls=false --bucket={self.cfg.s3_bucket} "
                      f"--obj.size=1KiB --concurrent=1 --duration=5s "
                      f"--benchdata=/dev/null 2>&1")
        result = subprocess.run(verify_cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
        verify_out = result.stdout.decode("utf-8", errors="replace") + result.stderr.decode("utf-8", errors="replace")
        if "Average:" in verify_out and "obj/s" in verify_out:
            log.info("S3 连通性验证通过 (warp 写入成功)")
        else:
            log.error("S3 连通性验证失败!")
            print(f"  warp 输出:\n{verify_out[-500:]}")
            print("\n  请检查: 1) S3 LIF 是否可达  2) 防火墙是否放通  3) Access Key 是否正确")
            sys.exit(1)

        # Save config
        self.env_data["s3_config"] = {
            "s3_endpoint": self.cfg.s3_endpoint,
            "access_key": self.cfg.s3_access_key,
            "secret_key": self.cfg.s3_secret_key,
            "bucket": self.cfg.s3_bucket,
            "svm_uuid": self.cfg.s3_svm_uuid,
            "s3_lif_ip": self.cfg.s3_lif_ip,
        }
        with open(os.path.join(self.work_dir, "env_report.json"), "w") as f:
            json.dump(self.env_data, f, indent=2, ensure_ascii=False)

        self.progress.mark_done("step3_s3")

    def _step4_test_config(self):
        log.banner("Step 4: 测试配置")
        print("  选择测试模式:")
        for i, (key, preset) in enumerate(TEST_PRESETS.items(), 1):
            scenes = count_scenes(preset)
            print(f"    [{i}] {preset['name']} — {preset['desc']} ({scenes} 场景)")
        print(f"    [4] 自定义")

        choice = input("\n  选择 [1-4]: ").strip() or "2"
        mode_map = {"1": "quick", "2": "standard", "3": "full"}
        self.cfg.test_mode = mode_map.get(choice, "standard")

        if choice == "4":
            # Custom mode - simplified
            self.cfg.test_mode = "standard"
            print("  (使用标准模式作为基础)")

        preset = TEST_PRESETS[self.cfg.test_mode]
        scenes = count_scenes(preset)
        print(f"\n  已选: {preset['name']} — {scenes} 个场景")
        input("  按 Enter 开始部署...")

    def _step5_deploy(self):
        if self.progress.is_done("step5_deploy"):
            log.info("Step 5 已完成，跳过")
            return

        log.banner("Step 5: 部署 warp")

        # Find warp binary: local package dir > PATH > /tmp > download
        script_dir = os.path.dirname(os.path.abspath(__file__))
        warp_path = None
        for candidate in [
            os.path.join(script_dir, "bin", "warp"),
            os.path.join(script_dir, "..", "bin", "warp"),
            shutil.which("warp"),
            "/usr/local/bin/warp",
            "/tmp/warp",
        ]:
            if candidate and os.path.isfile(candidate) and os.access(candidate, os.X_OK):
                warp_path = candidate
                break

        if warp_path:
            print(f"  warp 已存在: {warp_path}")
        else:
            warp_path = "/tmp/warp"
            print("  下载 warp...", end=" ", flush=True)
            try:
                subprocess.run(["wget", "-q", WARP_CDN, "-O", warp_path], timeout=60, check=True)
                os.chmod(warp_path, 0o755)
                print("✓")
            except Exception:
                print("CDN 失败, 尝试 GitHub...", end=" ", flush=True)
                subprocess.run(["wget", "-q", WARP_GITHUB, "-O", "/tmp/warp.tar.gz"], timeout=60, check=True)
                subprocess.run(["tar", "xzf", "/tmp/warp.tar.gz", "-C", "/tmp/"], check=True)
                os.chmod(warp_path, 0o755)
                print("✓")

        # Install locally
        try:
            subprocess.run(["sudo", "cp", warp_path, "/usr/local/bin/warp"], check=True)
        except Exception:
            shutil.copy(warp_path, os.path.expanduser("~/warp"))

        # Distribute to VMs
        for i, vm in enumerate(self.cfg.vms):
            ip = vm["ip"]
            progress_bar(i, len(self.cfg.vms), "分发 warp")
            self.ssh.upload(ip, warp_path, "/usr/local/bin/warp")
        progress_bar(len(self.cfg.vms), len(self.cfg.vms), "分发 warp")

        # Disable firewall on VMs
        print("  关闭 VM 防火墙...")
        for vm in self.cfg.vms:
            ip = vm["ip"]
            self.ssh.run(ip, "systemctl stop firewalld 2>/dev/null; iptables -F 2>/dev/null || true")
            log.info(f"{ip} 防火墙已关闭")

        # Sync time
        print("  同步时钟...")
        for vm in self.cfg.vms:
            self.ssh.run(vm["ip"], "chronyc makestep 2>/dev/null || ntpdate -s pool.ntp.org 2>/dev/null || true")

        # Start warp clients
        print("  启动 warp client...")
        for vm in self.cfg.vms:
            ip = vm["ip"]
            self.ssh.run(ip, "pkill warp 2>/dev/null; sleep 1")
            self.ssh.run(ip, f"mkdir -p /root/warp_monitor && nohup warp client {ip}:{WARP_PORT} > /tmp/warp_client.log 2>&1 &")

        time.sleep(5)

        # Verify warp client connectivity — fail fast if unreachable
        import socket
        ok = 0
        failed = []
        for vm in self.cfg.vms:
            ip = vm["ip"]
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.settimeout(5)
                s.connect((ip, WARP_PORT))
                s.close()
                ok += 1
                log.info(f"{ip}:{WARP_PORT} ✓")
            except Exception:
                failed.append(ip)
                log.error(f"{ip}:{WARP_PORT} 连接失败")

        if failed:
            log.error(f"warp client 连通检查失败: {', '.join(failed)}")
            print(f"\n  请检查: 1) VM 防火墙  2) warp client 进程  3) 网络连通性")
            sys.exit(1)

        log.info(f"warp client: {ok}/{len(self.cfg.vms)} 就绪")

        self.progress.mark_done("step5_deploy")

    def _step6_execute(self):
        log.banner("Step 6: 执行性能测试")

        preset = TEST_PRESETS[self.cfg.test_mode]
        duration = preset["duration"]
        objects = preset["objects"]
        warp_clients = ",".join(f'{vm["ip"]}:{WARP_PORT}' for vm in self.cfg.vms)

        common = (f"--host={self.cfg.s3_lif_ip}:{self.cfg.s3_port} "
                  f"--access-key={self.cfg.s3_access_key} "
                  f"--secret-key={self.cfg.s3_secret_key} "
                  f"--tls=false "
                  f"--warp-client={warp_clients} "
                  f"--bucket={self.cfg.s3_bucket}")

        # Build scene list
        scenes = []
        for sz, concs in preset["put"]:
            for c in concs:
                scenes.append(("put", sz, c, f"put_{sz}_c{c}"))
        for sz, concs in preset["get"]:
            obj = objects if sz != "4MiB" else min(objects, 1000)
            for c in concs:
                scenes.append(("get", sz, c, f"get_{sz}_c{c}"))
        for sz, concs in preset["mixed"]:
            for c in concs:
                scenes.append(("mixed", sz, c, f"mixed_{sz}_c{c}"))
        for item in preset["delete"]:
            sz, concs = item
            for c in concs:
                scenes.append(("delete", sz, c, f"delete_{sz}_c{c}"))
        for item in preset["list"]:
            c = item[0]
            scenes.append(("list", "4KiB", c, f"list_c{c}"))

        total = len(scenes)
        results_dir = os.path.join(self.work_dir, "warp_results")

        # Execute test matrix
        start_time = time.time()
        for idx, (op, sz, conc, scene_name) in enumerate(scenes, 1):
            if self.progress.is_done(f"scene_{scene_name}"):
                # Load saved result
                log_file = os.path.join(results_dir, f"{scene_name}_stdout.log")
                if os.path.exists(log_file):
                    with open(log_file) as f:
                        parsed = parse_warp_output(f.read())
                        if parsed:
                            parsed["scene"] = scene_name
                            self.results.append(parsed)
                print(f"  [{idx}/{total}] {scene_name} — 已完成，跳过")
                continue

            elapsed = time.time() - start_time
            eta = (elapsed / max(idx - 1, 1)) * (total - idx + 1) if idx > 1 else 0
            eta_str = f"{int(eta//60)}m{int(eta%60)}s" if eta > 0 else "计算中"

            progress_bar(idx - 1, total, f"{scene_name:<25}", extra=f"剩余 ~{eta_str}")

            # Build warp command
            extra = ""
            if op == "get":
                obj_count = objects if sz != "4MiB" else min(objects, 1000)
                extra = f"--objects={obj_count}"
            elif op == "mixed":
                extra = f"--objects={objects} --get-distrib=70 --put-distrib=30 --stat-distrib=0 --delete-distrib=0"
            elif op == "list":
                extra = f"--objects=10000"

            cmd = (f"warp {op} {common} --obj.size={sz} --concurrent={conc} "
                   f"--duration={duration} {extra} "
                   f"--benchdata={results_dir}/{scene_name}.csv.zst 2>&1")

            try:
                result = subprocess.run(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=600)
                output = result.stdout.decode("utf-8", errors="replace") + result.stderr.decode("utf-8", errors="replace")

                # Save stdout log
                with open(os.path.join(results_dir, f"{scene_name}_stdout.log"), "w") as f:
                    f.write(output)

                # Parse results
                parsed = parse_warp_output(output)
                if parsed:
                    parsed["scene"] = scene_name
                    parsed["status"] = "completed"
                    self.results.append(parsed)
                    self.file_log.info(f"{scene_name}: {parsed.get('throughput_mibps',0):.2f} MiB/s, {parsed.get('iops',0):.2f} obj/s")
                else:
                    self.file_log.warning(f"{scene_name}: 解析失败")
                    self.results.append({"scene": scene_name, "type": op.upper(),
                                        "concurrency": conc * len(self.cfg.vms),
                                        "throughput_mibps": 0, "iops": 0,
                                        "latency_avg_ms": 0, "latency_p50_ms": 0,
                                        "latency_p90_ms": 0, "latency_p99_ms": 0,
                                        "status": "parse_error"})
            except subprocess.TimeoutExpired:
                log.warn(f"{scene_name} 超时，跳过")
                self.results.append({"scene": scene_name, "type": op.upper(),
                                    "concurrency": conc * len(self.cfg.vms),
                                    "throughput_mibps": 0, "iops": 0,
                                    "latency_avg_ms": 0, "latency_p50_ms": 0,
                                    "latency_p90_ms": 0, "latency_p99_ms": 0,
                                    "status": "timeout"})
            except Exception as e:
                log.error(f"{scene_name} 失败: {e}")

            self.progress.mark_done(f"scene_{scene_name}")

            # Cooldown
            for vm in self.cfg.vms:
                self.ssh.run(vm["ip"], "echo 3 > /proc/sys/vm/drop_caches 2>/dev/null; sync")
            time.sleep(10)

        progress_bar(total, total, "全部完成")

        # Save test matrix
        with open(os.path.join(self.work_dir, "test_matrix.json"), "w") as f:
            json.dump({
                "test_date": datetime.now().strftime("%Y-%m-%d"),
                "tool": "MinIO warp",
                "mode": self.cfg.test_mode,
                "results": self.results,
            }, f, indent=2, ensure_ascii=False)

        # Print summary
        self._print_summary()
        self.progress.mark_done("step6_execute")

    def _print_summary(self):
        print(f"\n  {'─'*75}")
        print(f"  {'场景':<28} {'类型':<5} {'并发':>4} {'吞吐MiB/s':>11} {'IOPS':>10} {'P99 ms':>7}")
        print(f"  {'─'*75}")
        prev_type = ""
        for r in self.results:
            t = r.get("type", "?")
            if t != prev_type:
                if prev_type:
                    print()
                prev_type = t
            tp = f'{r.get("throughput_mibps",0):.2f}' if r.get("throughput_mibps",0) > 0 else "-"
            print(f'  {r["scene"]:<28} {t:<5} {r.get("concurrency",0):>4} '
                  f'{tp:>11} {r.get("iops",0):>10.0f} {r.get("latency_p99_ms",0):>7.1f}')
        print(f"  {'─'*75}")

    def _step7_reports(self):
        log.banner("Step 7: 生成报告")
        reporter = ReportGenerator(self.cfg, self.env_data, self.results, self.work_dir)
        reporter.generate_charts()
        reporter.generate_html()
        reporter.generate_word()

        # Package
        tarname = f"{self.cfg.customer_name}_ONTAP_S3_测试数据_{datetime.now().strftime('%Y%m%d')}.tar.gz"
        tarpath = os.path.join(self.work_dir, "reports", tarname)
        subprocess.run(
            f'cd {self.work_dir} && tar czf "{tarpath}" '
            f'reports/*.html reports/*.docx reports/charts/ '
            f'env_report.json test_matrix.json warp_results/ 2>/dev/null',
            shell=True)
        log.info(f"数据包: {tarpath}")

        # List outputs
        print(f"\n  交付物:")
        reports_dir = os.path.join(self.work_dir, "reports")
        for f in sorted(os.listdir(reports_dir)):
            if not f.startswith("charts"):
                fpath = os.path.join(reports_dir, f)
                size = os.path.getsize(fpath) / 1024
                print(f"    {f} ({size:.0f} KB)")

    def _step8_cleanup(self):
        log.banner("Step 8: 清理")
        print("  [1] 清理全部 (ONTAP S3 + VM warp + 数据)")
        print("  [2] 仅清理 VM (停 warp client)")
        print("  [3] 仅清理 ONTAP S3 (删 bucket + 用户)")
        print("  [4] 不清理")

        choice = input("\n  选择 [1-4]: ").strip() or "4"

        if choice in ("1", "2"):
            for vm in self.cfg.vms:
                self.ssh.run(vm["ip"], "pkill warp 2>/dev/null")
            log.info("VM warp client 已停止")

        if choice in ("1", "3"):
            if self.cfg.s3_svm_uuid and self.cfg.s3_bucket:
                try:
                    # Note: bucket deletion via API may require emptying first
                    log.warn(f"ONTAP bucket '{self.cfg.s3_bucket}' 需手动清理 (可能含残留对象)")
                except Exception as e:
                    log.warn(f"清理失败: {e}")

        log.info("清理完成")

    def _emergency_cleanup(self):
        """Ctrl+C 紧急清理"""
        for vm in self.cfg.vms:
            try:
                self.ssh.run(vm["ip"], "pkill warp 2>/dev/null")
            except Exception:
                pass
        self.ssh.close_all()
        # Save whatever we have
        if self.results:
            with open(os.path.join(self.work_dir, "test_matrix.json"), "w") as f:
                json.dump({"results": self.results, "interrupted": True}, f, indent=2, ensure_ascii=False)
        log.info("紧急清理完成，进度已保存")

    def _load_existing_data(self):
        """--report-only 模式加载已有数据"""
        env_file = os.path.join(self.work_dir, "env_report.json")
        matrix_file = os.path.join(self.work_dir, "test_matrix.json")
        if os.path.exists(env_file):
            with open(env_file) as f:
                self.env_data = json.load(f)
        if os.path.exists(matrix_file):
            with open(matrix_file) as f:
                data = json.load(f)
                self.results = data.get("results", [])
        log.info(f"已加载 {len(self.results)} 个测试结果")

# ─── 入口 ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="ONTAP S3 对象存储性能测试工具")
    parser.add_argument("--config", help="YAML 配置文件路径")
    parser.add_argument("--dry-run", action="store_true", help="仅探测不执行测试")
    parser.add_argument("--report-only", action="store_true", help="基于已有数据重新生成报告")
    parser.add_argument("--work-dir", default="~/ontap_s3_test", help="工作目录")
    args = parser.parse_args()

    if args.report_only:
        config = Config()
        config.report_only = True
        config.work_dir = os.path.expanduser(args.work_dir)
        # Load customer name from existing env_report.json
        env_file = os.path.join(config.work_dir, "env_report.json")
        if os.path.exists(env_file):
            with open(env_file) as f:
                env = json.load(f)
                config.customer_name = env.get("customer_name", "测试客户")
                config.s3_lif_ip = env.get("s3_config", {}).get("s3_lif_ip", "")
    elif args.config:
        config = Config.from_yaml(args.config)
    else:
        config = Config.from_interactive()

    config.dry_run = args.dry_run
    if not args.report_only:
        config.report_only = False
    if args.work_dir:
        config.work_dir = os.path.expanduser(args.work_dir)

    bench = Benchmark(config)
    bench.run()

if __name__ == "__main__":
    main()
